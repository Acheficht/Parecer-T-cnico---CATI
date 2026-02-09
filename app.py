import streamlit as st
import time # Importado apenas para simular o "Verificar Status"

# ==========================================
# 🛑 CONTROLE DE MANUTENÇÃO (LIGAR/DESLIGAR)
# ==========================================
EM_MANUTENCAO = False  # Mude para False quando quiser liberar o site normal

if EM_MANUTENCAO:
    # --- CONFIGURAÇÃO DA TELA DE MANUTENÇÃO ---
    st.set_page_config(page_title="Em Manutenção", page_icon="🚧", layout="centered")

    # Layout: Coluna da esquerda (Mensagem) e direita (Status)
    col1, col2 = st.columns([2, 1])

    with col1:
        st.title("🚧 Estamos em Manutenção")
        st.markdown("""
        ### O sistema está evoluindo.
        Estamos implementando uma atualização radical para a **Versão 2.0**.
        
        **O que muda?**
        * 🚀 Melhor performance.
        * 📂 Geração de documentos mais rápida.
        * ✨ Nova interface visual.
        
        Agradecemos a paciência!
        """)
        st.write("")
        st.info("Status atual: **Finalizando ajustes no servidor**")
        
        # Barra de progresso visual
        st.progress(85)

    with col2:
        st.subheader("🛠️ Ferramentas")
        st.write("Precisa de ajuda urgente?")
        
        # Botão simples de status (sem banco de dados)
        if st.button("🔄 Verificar Status"):
            with st.spinner("Checando servidores..."):
                time.sleep(1.5) # Simula um carregamento
                st.success("Sistemas: 🟢 Online")
                st.warning("App: 🟡 Em Atualização")
        
        st.divider()
        
        st.caption("Dúvidas ou suporte:")
        st.code("carlos.car.cati@gmail.com", language="text")

    st.divider()
    st.caption("© 2026 Equipe de Desenvolvimento - PRODESP")
    
    # 🛑 BLOQUEIO TOTAL
    # O comando abaixo impede que o resto do código rode.
    st.stop() 

# ==========================================
# 🚀 SEU CÓDIGO DO SITE NORMAL COMEÇA AQUI
# ==========================================
# Tudo abaixo desta linha só vai aparecer quando você mudar
# EM_MANUTENCAO = False lá no topo.

# ... Cole aqui suas 1000 linhas de código do sistema principal ...
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from fpdf import FPDF
import io
import json
import os
import tempfile
import base64
from datetime import date
from PIL import Image
import uuid
import copy
import random

# --- NOVA IMPORTAÇÃO NECESSÁRIA ---
# Certifique-se de instalar: pip install streamlit-pdf-viewer
try:
    from streamlit_pdf_viewer import pdf_viewer
except ImportError:
    st.error("Biblioteca 'streamlit-pdf-viewer' não instalada. Adicione ao requirements.txt")

# --- CONFIGURAÇÃO DO LOGO ---
ARQUIVO_LOGO = "logo.png.png"

# --- 1. CONFIGURAÇÃO INICIAL E CSS ---
# ATUALIZAÇÃO: Versão alterada para v2.0
st.set_page_config(page_title="Parecer Técnico v2.0 (Formatting Update)", layout="wide")

# --- SISTEMA DE FRASES E RARIDADE (MANTIDO IGUAL) ---
FRASES_SISTEMA = {
    "COMUM": {
        "chance": 0.485, 
        "cor": "#4CAF50",  
        "icone": "🟢",
        "classe": "loot-common",
        "frases": [
            "“Feito é melhor que perfeito.” — Sheryl Sandberg",
            "“Comece antes de estar pronto.” — Steven Pressfield",
            "“Um passo pequeno ainda é progresso.”",
            "“Consistência vence talento desorganizado.”",
            "“Disciplina é liberdade.” — Jocko Willink",
            "“Você não precisa de sorte, precisa de ação.”",
            "“Hoje é um ótimo dia para não desistir.”",
            "“Treino ruim ainda vence dia sem treino.”",
            "“Foco é dizer não.” — Steve Jobs",
            "“O básico bem feito é poderoso.”",
            "“Erro rápido, progresso rápido.”",
            "“Pressa passa. Resultado fica.”",
            "“Se começou, continue.”",
            "“Melhor cansado do que arrependido.”",
            "“Motivação falha — rotina não.”",
            "“Você não travou — você está aprendendo.”",
            "“Só termina quem continua.”",
            "“Não complique o que funciona.”",
            "“Progresso > perfeição.”",
            "“A prática resolve dúvidas.”",
            "“Se organize menos — comece mais.”",
            "“Sem drama, mais ação.”",
            "“Passos simples constroem coisas grandes.”",
            "“Constância é um superpoder.”",
            "“Café, foco e execução.”"
        ]
    },
    "INCOMUM": {
        "chance": 0.25, 
        "cor": "#2196F3", 
        "icone": "🔹",
        "classe": "loot-uncommon",
        "frases": [
            "“A ação é a chave fundamental para todo sucesso.” — Pablo Picasso",
            "“Quem tem um porquê enfrenta quase qualquer como.” — Nietzsche",
            "“Você é aquilo que faz repetidamente.” — Aristóteles",
            "“Coragem não é ausência de medo, é decisão apesar dele.”",
            "“Grandes resultados exigem desconforto.”",
            "“Não espere motivação — construa hábito.”",
            "“O esforço de hoje é o respeito de amanhã.”",
            "“O que você faz em silêncio constrói seu futuro.”",
            "“Ideias valem pouco sem execução.”",
            "“Desconforto é taxa de crescimento.”",
            "“Rotina forte vence emoção fraca.”",
            "“Você não precisa sentir — precisa fazer.”"
        ]
    },
    "RARA": {
        "chance": 0.15, 
        "cor": "#FF9800", 
        "icone": "🔸",
        "classe": "loot-rare",
        "frases": [
            "“No meio do inverno aprendi que havia em mim um verão invencível.” — Albert Camus",
            "“Sofremos mais na imaginação do que na realidade.” — Sêneca",
            "“A vida encolhe ou expande na proporção da coragem.” — Anaïs Nin",
            "“Quem olha para fora sonha; quem olha para dentro desperta.” — Carl Jung",
            "“Torne-se quem você é.” — Nietzsche",
            "“O onde há dificuldade, há caminho.”"
        ]
    },
    "EPICA": {
        "chance": 0.05, 
        "cor": "#9C27B0", 
        "icone": "🔶",
        "classe": "loot-epic",
        "frases": [
            "“Aquele que tem um motivo para viver pode suportar quase qualquer coisa.” — Nietzsche",
            "“A ferida é o lugar por onde a luz entra.” — Rumi",
            "“Você não é uma gota no oceano — é o oceano inteiro em uma gota.” — Rumi",
            "“O homem é condenado a ser livre.” — Jean-Paul Sartre"
        ]
    },
    "LENDARIA": {
        "chance": 0.03, 
        "cor": "#00BCD4", 
        "icone": "💎",
        "classe": "loot-legendary",
        "frases": [
            "“Não é que tenhamos pouco tempo — desperdiçamos muito.” — Sêneca",
            "“Entre o estímulo e a resposta existe um espaço — nele está sua liberdade.” — Viktor Frankl"
        ]
    },
    "DIVINA": {
        "chance": 0.005, 
        "cor": "#FFD700", 
        "icone": "✨",
        "classe": "loot-god",
        "frases": [
            "“Quem encontra propósito transforma sofrimento em combustível.”"
        ]
    }
}

# CSS para as animações de raridade
st.markdown("""
<style>
    @keyframes shine {
        0% {background-position: -100px;}
        100% {background-position: 300px;}
    }
    @keyframes pulse {
        0% { transform: scale(1); }
        50% { transform: scale(1.02); }
        100% { transform: scale(1); }
    }
    @keyframes rainbow { 
        0%{background-position:0% 50%}
        50%{background-position:100% 50%}
        100%{background-position:0% 50%}
    }
    
    .loot-box {
        padding: 20px;
        border-radius: 10px;
        text-align: center;
        margin-bottom: 20px;
        color: #333;
        font-family: 'Arial', sans-serif;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }

    .msg-number {
        font-size: 0.8em;
        color: #666;
        margin-top: 5px;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    /* COMUM */
    .loot-common {
        border: 2px solid #4CAF50;
        background-color: #E8F5E9;
    }
    
    /* INCOMUM */
    .loot-uncommon {
        border: 2px solid #2196F3;
        background-color: #E3F2FD;
        box-shadow: 0 0 10px rgba(33, 150, 243, 0.3);
    }
    
    /* RARA */
    .loot-rare {
        border: 2px solid #FF9800;
        background: linear-gradient(135deg, #FFF3E0 0%, #FFE0B2 100%);
        box-shadow: 0 0 15px rgba(255, 152, 0, 0.4);
        animation: pulse 3s infinite;
    }
    
    /* EPICA */
    .loot-epic {
        border: 2px solid #9C27B0;
        background-color: #F3E5F5;
        box-shadow: 0 0 20px rgba(156, 39, 176, 0.5);
        border-radius: 12px;
        animation: pulse 2s infinite;
    }
    
    /* LENDARIA */
    .loot-legendary {
        border: 3px solid #00BCD4;
        background: linear-gradient(45deg, #E0F7FA, #FFFFFF, #E0F7FA);
        box-shadow: 0 0 25px rgba(0, 188, 212, 0.7);
        position: relative;
        overflow: hidden;
    }
    .loot-legendary::before {
        content: '';
        position: absolute;
        top: 0; left: -150px;
        width: 100px; height: 100%;
        background: rgba(255,255,255,0.6);
        transform: skewX(-25deg);
        animation: shine 3s infinite;
    }
    
    /* DIVINA */
    .loot-god {
        border: 4px solid transparent;
        background: linear-gradient(white, white) padding-box,
                  linear-gradient(45deg, #FF0000, #FF7300, #FFFB00, #48FF00, #00FFD5, #002BFF, #7A00FF, #FF00C8, #FF0000) border-box;
        box-shadow: 0 0 40px rgba(255, 215, 0, 0.8);
        animation: pulse 1.5s infinite;
        color: black;
        font-weight: bold;
    }
    
    .quote-text {
        font-size: 1.3em;
        font-style: italic;
        margin: 15px 0;
    }
    .rarity-label {
        font-weight: bold;
        text-transform: uppercase;
        font-size: 0.9em;
        letter-spacing: 2px;
    }
</style>
""", unsafe_allow_html=True)

# --- MENSAGEM DE ATUALIZAÇÃO (POP-UP) ---
@st.dialog("✨ Dados da Atualização")
def mostrar_novidades():
    st.markdown("""
    Bem-vindo à nova versão! Confira as melhorias implementadas para otimizar seu trabalho:

    **🚀 Novas Funcionalidades**
    1.  **Reordenação de Itens:** Agora é possível mover as caixas de inconsistências para cima ⬆️ ou para baixo ⬇️ conforme sua necessidade.
    2.  **Histórico de Ações (Undo/Redo):** Cometeu um erro? Utilize os botões de "Voltar" e "Avançar" (estilo Ctrl+Z) para desfazer ou refazer alterações.
    3.  **Nomes Personalizados:** Agora você tem total liberdade para nomear os arquivos antes de baixar (PDF, Word) e também ao salvar seus backups.
    4.  **Identidade Visual:** A logo do CAR foi integrada ao cabeçalho do site e documentos.
    5.  **Formatação Avançada:** Agora você pode usar **Negrito**, *Itálico* e __Sublinhado__ nos textos!
        * Use `**texto**` para **negrito**.
        * Use `*texto*` para *itálico*.
        * Use `__texto__` (dois underlines) para sublinhado.
        * **Dica:** Você pode combiná-los! Ex: `**__Negrito e Sublinhado__**`.
    6.  **Mensagem do Dia:** Documento finalizado! Gere o PDF ou Word e receba sua mensagem do dia. Será que hoje vem uma classe **DIVINA**? Tente sua SORTE!

    **🔧 Correção de Bugs e Ajustes**
    * **Nomenclatura:** Alterado o nome de "Gerador de Parecer Técnico (Com Item Personalizável)" para "Justificativa do Parecer Técnico".
    * **Visualização PDF:** Corrigido bug que impedia a visualização no Chrome/Edge.
    * **Imagens:** Agora as imagens ocupam metade do tamanho original e são centralizadas automaticamente.
    * **Numeração de Página:** Reposicionada para o canto inferior direito (rodapé).
    * **Formatação Word:** Remoção de espaçamentos excessivos.
    
    ---
    *OBS: Qualquer dúvida, reclamação ou sugestão, contate o suporte por meio do e-mail:* <span style="color: blue; text-decoration: underline;">carlos.car.cati@gmail.com</span>
    """, unsafe_allow_html=True)
    
    if st.button("Entendi, vamos trabalhar!"):
        st.rerun()

# Lógica para mostrar o pop-up apenas na primeira vez
if "modal_visto" not in st.session_state:
    mostrar_novidades()
    st.session_state["modal_visto"] = True

# --- CSS GERAL DO SITE ---
st.markdown("""
<style>
    div[data-testid="stVerticalBlock"] > div > button {
        border-color: #ff4b4b;
        color: #ff4b4b;
    }
    div[data-testid="stVerticalBlock"] > div > button:hover {
        background-color: #ff4b4b;
        color: white;
    }
    .stButton > button[kind="primary"] {
        background-color: #4CAF50;
        color: white;
        border: none;
    }
    .img-container {
        border: 1px solid #ddd;
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 10px;
        text-align: center;
        background-color: #f9f9f9;
    }
    .block-container { padding-top: 2rem; }
    
    button[kind="secondary"] {
        padding: 0.25rem 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# --- 2. LISTA MESTRA DE OPÇÕES ---
OPCOES_LISTA = [
  "Inconsistências em Ficha do imóvel",
  "Inconsistências em Sobreposição com outros IRs",
  "Inconsistência em Sobreposição com Unidade de Conservação de Uso Sustentável",
  "Outras Sobreposições",
  "Inconsistências em Áreas embargadas",
  "Inconsistências em Assentamentos",
  "Inconsistências em UC",
  "Inconsistências em Cobertura do solo", 
  "Inconsistências em Infraestrutura e utilidade pública",
  "Inconsistências em Reservatório para abastecimento ou geração de energia",
  "Inconsistências em APP hidrografia",
  "Inconsistências em APP Relevo",
  "Inconsistências em Uso restrito",
  "Inconsistências em outras APPs",
  "Inconsistências em RL averbada, RL aprovada e não averbada",
  "Inconsistências em Área de RL exigida por lei",
  "Inconsistências em Localização e cobertura do solo",
  "Inconsistências em Regularidade do IR",
  "Inconsistência Adicional",
  "Observação",
  "Item Personalizado ✏️"
]

# --- 3. GERENCIAMENTO DE ESTADO E HISTÓRICO ---
if 'dados' not in st.session_state:
  st.session_state['dados'] = {
      "car": "", "sp_not": "", "imovel": "", 
      "nome": "", "doc": "", "cidade": "Mogi das Cruzes",
      "itens": [], 
      "textos": {}, 
      "imagens_b64": {} 
    }

if 'uploader_ids' not in st.session_state:
  st.session_state['uploader_ids'] = {}

if 'historico_undo' not in st.session_state:
  st.session_state['historico_undo'] = [] 
if 'historico_redo' not in st.session_state:
  st.session_state['historico_redo'] = [] 

if 'frase_atual_download' not in st.session_state:
    st.session_state['frase_atual_download'] = None

# --- 4. FUNÇÕES DE HISTÓRICO (UNDO/REDO) ---

def salvar_estado_no_historico():
  estado_atual = copy.deepcopy(st.session_state['dados'])
  st.session_state['historico_undo'].append(estado_atual)
  st.session_state['historico_redo'] = []

def desfazer_acao():
    if st.session_state['historico_undo']:
        estado_atual = copy.deepcopy(st.session_state['dados'])
        st.session_state['historico_redo'].append(estado_atual)
        
        estado_anterior = st.session_state['historico_undo'].pop()
        st.session_state['dados'] = estado_anterior
        
        keys_map = ["car","sp_not","imovel","nome", "doc", "cidade"]
        for k in keys_map:
          st.session_state[f"input_{k}"] = estado_anterior.get(k, "")
        st.toast("⏪ Ação desfeita!")

def refazer_acao():
    if st.session_state['historico_redo']:
        estado_atual = copy.deepcopy(st.session_state['dados'])
        st.session_state['historico_undo'].append(estado_atual)
        
        estado_futuro = st.session_state['historico_redo'].pop()
        st.session_state['dados'] = estado_futuro
        
        keys_map = ["car","sp_not","imovel","nome", "doc", "cidade"]
        for k in keys_map:
          st.session_state[f"input_{k}"] = estado_futuro.get(k, "")
        st.toast("↪️ Ação refeita!")

# --- 5. FUNÇÕES DE MANIPULAÇÃO DE ITENS ---

def adicionar_item():
  salvar_estado_no_historico()
  opcao = st.session_state.get("selecao_adicionar")
  if opcao:
      novo_id = str(uuid.uuid4())
      eh_personalizado = (opcao == "Item Personalizado ✏️")
      titulo_inicial = "" if eh_personalizado else opcao
      
      novo_item = {
          "id": novo_id, 
          "titulo": titulo_inicial, 
          "custom": eh_personalizado
      }
      st.session_state['dados']['itens'].append(novo_item)

def remover_item(id_item):
  salvar_estado_no_historico()
  st.session_state['dados']['itens'] = [
      item for item in st.session_state['dados']['itens'] if item['id'] != id_item
  ]
  if id_item in st.session_state['dados']['textos']:
      del st.session_state['dados']['textos'][id_item]
  if id_item in st.session_state['dados']['imagens_b64']:
      del st.session_state['dados']['imagens_b64'][id_item]

def mover_item(index, direcao):
  salvar_estado_no_historico()
  lista = st.session_state['dados']['itens']
  
  if direcao == 'cima' and index > 0:
      lista[index], lista[index-1] = lista[index-1], lista[index]
  elif direcao == 'baixo' and index < len(lista) - 1:
      lista[index], lista[index+1] = lista[index+1], lista[index]

def processar_upload(id_item):
  uid = st.session_state['uploader_ids'].get(id_item, 0)
  key_widget = f"uploader_{id_item}_{uid}"
  uploaded_file = st.session_state.get(key_widget)
  if uploaded_file:
      try:
          uploaded_file.seek(0)
          bytes_data = uploaded_file.read()
          b64_str = base64.b64encode(bytes_data).decode('utf-8')
          if id_item not in st.session_state['dados']['imagens_b64']:
             st.session_state['dados']['imagens_b64'][id_item] = []
          st.session_state['dados']['imagens_b64'][id_item].append(b64_str)
          st.session_state['uploader_ids'][id_item] = uid + 1
      except Exception as e:
          st.error(f"Erro no processamento: {e}")

# --- 6. FUNÇÕES AUXILIARES ---

def obter_data_extenso():
  meses = {1:'janeiro', 2:'fevereiro', 3:'março', 4:'abril', 5:'maio', 6:'junho', 7:'julho', 8:'agosto', 9:'setembro', 10:'outubro', 11:'novembro', 12:'dezembro'}
  hj = date.today()
  return f"{hj.day} de {meses[hj.month]} de {hj.year}"

def limpar_tudo():
  salvar_estado_no_historico()
  st.session_state['dados'] = {
      "car": "", "sp_not": "", "imovel": "", 
      "nome": "", "doc": "", "cidade": "", 
      "itens": [], "textos": {}, "imagens_b64": {}
    }
  st.session_state['uploader_ids'] = {}
  campos = ["car","sp_not","imovel","nome", "doc", "cidade"]
  for c in campos:
      st.session_state[f"input_{c}"] = ""

def limpar_campo_cabecalho(key_suffix):
  st.session_state[f"input_{key_suffix}"] = ""

def remover_imagem_especifica(id_item, index):
  salvar_estado_no_historico()
  if id_item in st.session_state['dados']['imagens_b64']:
      lista = st.session_state['dados']['imagens_b64'][id_item]
      if 0 <= index < len(lista):
          lista.pop(index)
          if not lista:
              del st.session_state['dados']['imagens_b64'][id_item]

def formatar_documento(n):
  if not n: return ""
  n = "".join(filter(str.isdigit, str(n)))
  if len(n) == 11: return f"{n[:3]}.{n[3:6]}.{n[6:9]}-{n[9:]}"
  elif len(n) == 14: return f"{n[:2]}.{n[2:5]}.{n[5:8]}/{n[8:12]}-{n[12:]}"
  return n

def b64_para_tempfile(b64_str):
  try:
      bytes_data = base64.b64decode(b64_str)
      file_obj = io.BytesIO(bytes_data)
      img = Image.open(file_obj)
      if img.mode in ('RGBA', 'LA'):
          background = Image.new(img.mode[:-1], img.size, (255, 255, 255))
          background.paste(img, img.split()[-1])
          img = background
      img = img.convert('RGB')
      with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tf:
          img.save(tf, format='JPEG', quality=95)
          return tf.name
  except Exception:
      return None

# --- PARSER DE FORMATAÇÃO ---
def parse_styled_text(text):
    segments = []
    current_text = []
    bold = False
    italic = False
    underline = False
    i = 0
    n = len(text)
    
    while i < n:
        if i + 1 < n and text[i:i+2] == '**':
             if current_text:
                 segments.append({"text": "".join(current_text), "bold": bold, "italic": italic, "underline": underline})
                 current_text = []
             bold = not bold
             i += 2
        elif i + 1 < n and text[i:i+2] == '__':
             if current_text:
                 segments.append({"text": "".join(current_text), "bold": bold, "italic": italic, "underline": underline})
                 current_text = []
             underline = not underline
             i += 2
        elif text[i] == '*':
             if current_text:
                 segments.append({"text": "".join(current_text), "bold": bold, "italic": italic, "underline": underline})
                 current_text = []
             italic = not italic
             i += 1
        else:
             current_text.append(text[i])
             i += 1
           
    if current_text:
        segments.append({"text": "".join(current_text), "bold": bold, "italic": italic, "underline": underline})
        
    return segments

# --- FUNÇÃO AUXILIAR PARA NÚMERO DE PÁGINA NO WORD ---
def criar_elemento_pagina(run):
  fldChar1 = OxmlElement('w:fldChar')
  fldChar1.set(ns.qn('w:fldCharType'), 'begin')

  instrText = OxmlElement('w:instrText')
  instrText.set(ns.qn('xml:space'), 'preserve')
  instrText.text = "PAGE"

  fldChar2 = OxmlElement('w:fldChar')
  fldChar2.set(ns.qn('w:fldCharType'), 'end')

  run._r.append(fldChar1)
  run._r.append(instrText)
  run._r.append(fldChar2)

# --- 7. GERAÇÃO DE PDF ---
class PDF(FPDF):
  def header(self):
      if os.path.exists(ARQUIVO_LOGO):
          self.image(ARQUIVO_LOGO, x=12.5, y=10, w=15.9, h=15.9)
      self.ln(20)

  def footer(self):
      self.set_y(-15)
      self.set_font('Arial', '', 10)
      self.cell(0, 10, str(self.page_no()), 0, 0, 'R')

  def write_markdown(self, text, line_height=7):
      text = str(text) if text else ""
      text = text.encode('latin-1', 'replace').decode('latin-1')
      linhas = text.split('\n')
      
      for linha in linhas:
          if not linha.strip():
              self.ln(line_height)
              continue
          
          self.write(line_height, "    ") # Indentação
          
          segmentos = parse_styled_text(linha)
          for seg in segmentos:
              style = ''
              if seg['bold']: style += 'B'
              if seg['italic']: style += 'I'
              if seg['underline']: style += 'U'
              
              self.set_font("Arial", style, 12)
              self.write(line_height, seg['text'])
          
          self.ln(line_height)

def gerar_pdf_bytes():
  try:
      pdf = PDF()
      pdf.set_margins(30, 30, 20)
      pdf.add_page()
      def safe_text(text): return text.encode('latin-1', 'replace').decode('latin-1') if text else ""

      pdf.set_font("Arial", 'B', 14); pdf.cell(0, 8, safe_text("Justificativa do Parecer Técnico"), ln=True, align='C')
      pdf.set_font("Arial", 'B', 12); pdf.cell(0, 6, safe_text(f"CAR: {st.session_state['dados']['car']}"), ln=True, align='C')
      pdf.cell(0, 6, safe_text(f"SP-NOT: {st.session_state['dados']['sp_not']}"), ln=True, align='C'); pdf.ln(5)
      
      pdf.set_font("Arial", 'B', 12); pdf.write(6, safe_text("Nome do Imóvel Rural: ")); pdf.set_font("Arial", '', 12); pdf.write(6, safe_text(st.session_state['dados']['imovel'])); pdf.ln(6)
      pdf.set_font("Arial", 'B', 12); pdf.write(6, safe_text("Nome: ")); pdf.set_font("Arial", '', 12); pdf.write(6, safe_text(st.session_state['dados']['nome'])); pdf.ln(6)
      pdf.set_font("Arial", 'B', 12); pdf.write(6, safe_text("CPF/CNPJ: ")); pdf.set_font("Arial", '', 12); pdf.write(6, safe_text(formatar_documento(st.session_state['dados']['doc']))); pdf.ln(10)
      
      for i, item_obj in enumerate(st.session_state['dados']['itens']):
          titulo = item_obj['titulo']
          if not titulo: titulo = "Sem Título Definido"
          id_item = item_obj['id']
          pdf.set_font("Arial", 'B', 12); pdf.cell(0, 8, safe_text(f"{i+1}. {titulo}"), ln=True)
          
          texto_raw = st.session_state['dados']['textos'].get(id_item, "")
          widget_val = st.session_state.get(f"txt_area_{id_item}")
          if widget_val is not None: texto_raw = widget_val
          
          pdf.write_markdown(texto_raw)
          
          lista_imgs = st.session_state['dados']['imagens_b64'].get(id_item, [])
          if lista_imgs:
              pdf.ln(2)
              for b64_img in lista_imgs:
                  temp_path = b64_para_tempfile(b64_img)
                  if temp_path:
                      try:
                          x_pos = (210 - 70) / 2
                          if pdf.get_y() > 220: pdf.add_page()
                          pdf.image(temp_path, x=x_pos, w=70)
                          pdf.ln(5)
                      finally:
                          if os.path.exists(temp_path): os.unlink(temp_path)
              pdf.ln(2)
          pdf.ln(4)
          
      pdf.ln(10); pdf.cell(0, 6, "________________________________________", ln=True, align='R')
      pdf.cell(0, 6, safe_text(st.session_state['dados']['nome']), ln=True, align='R')
      cidade_doc = st.session_state['dados']['cidade'] if st.session_state['dados']['cidade'] else "Mogi das Cruzes"
      pdf.cell(0, 6, safe_text(f"{cidade_doc}, {obter_data_extenso()}."), ln=True, align='R')
      return pdf.output(dest='S').encode('latin-1', 'replace')
  except Exception as e: return None

# --- 8. BARRA LATERAL ---
with st.sidebar:
  st.header("🎮 Controles")
  col_undo, col_redo = st.columns(2)
  
  tem_undo = len(st.session_state['historico_undo']) > 0
  tem_redo = len(st.session_state['historico_redo']) > 0
  
  with col_undo:
      if st.button("↩️ Voltar", disabled=not tem_undo, help="Desfazer a última ação crítica"):
          desfazer_acao()
          st.rerun()

  with col_redo:
      if st.button("↪️ Avançar", disabled=not tem_redo, help="Refazer a ação desfeita"):
          refazer_acao()
          st.rerun()
  
  st.write(f"<small style='color:gray'>Histórico: {len(st.session_state['historico_undo'])} passos</small>", unsafe_allow_html=True)
  st.markdown("---")

  st.header("🗂️ Arquivos")
  arquivo_upload = st.file_uploader("📂 Carregar Trabalho (.json)", type=["json"])
  
  if arquivo_upload is not None:
      if st.button("🔄 Confirmar Carregamento"):
          try:
              salvar_estado_no_historico()
              dados_carregados = json.load(arquivo_upload)
              
              if "selecionados" in dados_carregados and "itens" not in dados_carregados:
                  novos_itens = []
                  novos_textos = {}
                  novas_imgs = {}
                  for nome_antigo in dados_carregados.get("selecionados", []):
                      novo_id = str(uuid.uuid4())
                      novos_itens.append({"id": novo_id, "titulo": nome_antigo, "custom": False})
                      if nome_antigo in dados_carregados.get("textos", {}):
                          novos_textos[novo_id] = dados_carregados["textos"][nome_antigo]
                      img_val = dados_carregados.get("imagens_b64", {}).get(nome_antigo, [])
                      if isinstance(img_val, str): img_val = [img_val]
                      if img_val: novas_imgs[novo_id] = img_val
                  
                  dados_carregados["itens"] = novos_itens
                  dados_carregados["textos"] = novos_textos
                  dados_carregados["imagens_b64"] = novas_imgs

              st.session_state['dados'] = dados_carregados
              for k in ["car", "sp_not", "imovel", "nome", "doc", "cidade"]:
                  st.session_state[f"input_{k}"] = dados_carregados.get(k, "")
              
              st.session_state['uploader_ids'] = {} 
              st.success("✅ Carregado com Sucesso!")
              st.rerun()
          except Exception as e:
              st.error(f"Erro: {e}")
  
  st.markdown("---")
  st.subheader("💾 Salvar Backup")
  
  sugestao_backup = st.session_state['dados']['nome'].strip() or "backup_dados"
  sugestao_backup = "".join([c for c in sugestao_backup if c.isalnum() or c in (' ','-','_')]).strip()
  
  nome_backup = st.text_input("Nome do arquivo .json:", value=f"{sugestao_backup}.json")
  if not nome_backup.lower().endswith(".json"): nome_backup += ".json"

  dados_download = json.dumps(st.session_state['dados'], indent=4)
  st.download_button("💾 Baixar Backup", dados_download, nome_backup, "application/json")
  
  st.markdown("---")
  st.button("🗑️ Limpar Tudo", on_click=limpar_tudo, type="primary")

# --- 9. INTERFACE PRINCIPAL ---

# ATUALIZAÇÃO: Ajuste na proporção das colunas para aproximar Logo e Título
col_logo, col_titulo = st.columns([0.08, 0.92]) 
with col_logo:
  if os.path.exists(ARQUIVO_LOGO):
      st.image(ARQUIVO_LOGO, width=100)
with col_titulo:
  # ATUALIZAÇÃO: Nome do app alterado
  st.title("Justificativa do Parecer Técnico")

tab_edit, tab_preview = st.tabs(["✍️ Edição", "👁️ Pré-visualização Real (PDF)"])

with tab_edit:
  st.subheader("1. Cabeçalho")
  def campo_com_lixeira(label, key_suffix):
      c_input, c_btn = st.columns([0.85, 0.15])
      with c_input:
          val = st.text_input(label, key=f"input_{key_suffix}")
          st.session_state['dados'][key_suffix] = val
      with c_btn:
          st.write(""); st.write("")
          st.button("🗑️", key=f"del_header_{key_suffix}", on_click=limpar_campo_cabecalho, args=(key_suffix,))

  c1, c2 = st.columns(2)
  with c1:
      campo_com_lixeira("CAR:", "car")
      campo_com_lixeira("SP-NOT (Número):", "sp_not")
      campo_com_lixeira("Nome do Imóvel:", "imovel")
  with c2:
      campo_com_lixeira("Nome do Requerente:", "nome")
      campo_com_lixeira("CPF/CNPJ:", "doc")
      campo_com_lixeira("Cidade:", "cidade")

  st.markdown("---")
  st.subheader("2. Adicionar Inconsistências")
  
  col_sel, col_add = st.columns([0.8, 0.2])
  with col_sel:
      st.selectbox("Escolha o tipo de item para adicionar:", OPCOES_LISTA, key="selecao_adicionar")
  with col_add:
      st.write(""); st.write("")
      st.button("➕ Adicionar", on_click=adicionar_item, type="primary")

  st.markdown("---")
  
  lista_itens = st.session_state['dados']['itens']
  
  if not lista_itens:
      st.info("Nenhum item adicionado ainda. Use a caixa acima para começar.")
  else:
      st.markdown("""
        > 💡 **Dica:** Você pode reordenar os itens usando as setas para cima e para baixo.
        > **Formatação:** Use `**negrito**`, `*itálico*` e `__sublinhado__`.
      """)
      st.write(f"**Itens no Relatório:** {len(lista_itens)}")
      
      for i, item_obj in enumerate(lista_itens):
          titulo = item_obj['titulo']
          id_item = item_obj['id']
          eh_custom = item_obj.get('custom', False)
          numero = i + 1
          
          c_titulo, c_up, c_down, c_remove = st.columns([0.80, 0.05, 0.05, 0.10])
          
          with c_up:
             st.write("")
             if st.button("⬆️", key=f"up_{id_item}", disabled=(i==0), on_click=mover_item, args=(i, 'cima')):
                 pass

          with c_down:
             st.write("")
             if st.button("⬇️", key=f"down_{id_item}", disabled=(i==len(lista_itens)-1), on_click=mover_item, args=(i, 'baixo')):
                 pass

          with c_remove:
             st.write("")
             st.button("🗑️", key=f"btn_del_{id_item}", on_click=remover_item, args=(id_item,), help="Excluir este item")
          
          with c_titulo:
              display_title = titulo if titulo else "(Digite o título do item abaixo...)"
              with st.expander(f"**{numero}. {display_title}**", expanded=False):
                  
                  if eh_custom:
                      st.markdown("#### ✏️ Nome do Item Personalizado")
                      novo_titulo = st.text_input("Digite o título deste item:", value=titulo, key=f"titulo_custom_{id_item}")
                      if novo_titulo != titulo:
                          item_obj['titulo'] = novo_titulo

                  chave_txt = f"txt_area_{id_item}"
                  val_inicial = st.session_state['dados']['textos'].get(id_item, "")
                  texto = st.text_area("Descrição:", value=val_inicial, key=chave_txt, height=150, placeholder="Ex: **A)** Inconsistência na área X, com trechos em *itálico* ou __sublinhado__.")
                  
                  if texto: st.session_state['dados']['textos'][id_item] = texto
                  elif id_item in st.session_state['dados']['textos'] and not texto:
                       del st.session_state['dados']['textos'][id_item]
                  
                  st.markdown("#### 🖼️ Imagens")
                  lista_imgs = st.session_state['dados']['imagens_b64'].get(id_item, [])
                  
                  if lista_imgs:
                      cols_imgs = st.columns(3)
                      for idx_img, img_b64 in enumerate(lista_imgs):
                          with cols_imgs[idx_img % 3]:
                              st.markdown("<div class='img-container'>", unsafe_allow_html=True)
                              st.image(base64.b64decode(img_b64), width="stretch")
                              if st.button("❌", key=f"del_img_{id_item}_{idx_img}"):
                                   remover_imagem_especifica(id_item, idx_img)
                                   st.rerun()
                              st.markdown("</div>", unsafe_allow_html=True)
                  
                  uid = st.session_state['uploader_ids'].get(id_item, 0)
                  key_uploader = f"uploader_{id_item}_{uid}"
                  st.file_uploader(f"Adicionar imagem", type=['png','jpg','jpeg'], key=key_uploader, on_change=processar_upload, args=(id_item,), label_visibility="collapsed")

with tab_preview:
  st.info("Visualização do PDF.")
  pdf_bytes_preview = gerar_pdf_bytes()
  
  if pdf_bytes_preview:
      # --- CORREÇÃO AQUI: USANDO A BIBLIOTECA DEDICADA ---
      # Não usamos mais base64 + iframe manual
      try:
          pdf_viewer(input=pdf_bytes_preview, width=700, height=800)
      except Exception as e:
          st.error(f"Erro ao renderizar PDF. Verifique se 'streamlit-pdf-viewer' está instalado. Detalhe: {e}")
  else:
      st.warning("Preencha os dados.")

# --- LÓGICA DE GERAÇÃO DA FRASE (SEPARADA DO DISPLAY) ---
def gerar_nova_frase_motivacional():
    # 1. Contar total de frases e criar mapa de índice
    todas_as_frases_flat = []
    # Iterar pelas raridades para criar uma lista plana única
    for tipo, dados in FRASES_SISTEMA.items():
        for f in dados['frases']:
            todas_as_frases_flat.append(f)
    
    total_mensagens = len(todas_as_frases_flat)

    # 2. Sortear a raridade
    tipos = list(FRASES_SISTEMA.keys())
    pesos = [FRASES_SISTEMA[t]["chance"] for t in tipos]
    raridade_sorteada = random.choices(tipos, weights=pesos, k=1)[0]
    
    # 3. Sortear a frase dentro da raridade
    dado_raridade = FRASES_SISTEMA[raridade_sorteada]
    frase_final = random.choice(dado_raridade["frases"])
    
    # 4. Encontrar o número desta mensagem (índice global)
    # Adicionamos +1 para ser de 1 a 50, não 0 a 49
    try:
        numero_mensagem = todas_as_frases_flat.index(frase_final) + 1
    except ValueError:
        numero_mensagem = "?"

    # 5. Salvar no session state
    st.session_state['frase_atual_download'] = {
        "frase": frase_final,
        "raridade": raridade_sorteada,
        "cor": dado_raridade['cor'],
        "icone": dado_raridade['icone'],
        "classe": dado_raridade['classe'],
        "numero": numero_mensagem,
        "total": total_mensagens
    }

# --- DIALOGO DE DOWNLOAD (MODIFICADO) ---
@st.dialog("✨ Momento de Inspiração")
def mostrar_mensagem_download(tipo_arquivo, dados_arquivo, nome_arquivo, mime_type):
    # Recupera a frase já sorteada (NÃO sorteia de novo)
    dados_frase = st.session_state.get('frase_atual_download')
    
    if dados_frase:
        # Exibir a "Loot Box" com HTML/CSS e NUMERAÇÃO
        st.markdown(f"""
        <div class="loot-box {dados_frase['classe']}">
            <div style="font-size: 3em;">{dados_frase['icone']}</div>
            <div class="rarity-label" style="color: {dados_frase['cor']};">{dados_frase['raridade']}</div>
            <div class="quote-text">
               {dados_frase['frase']}
            </div>
            <div class="msg-number">
                 Mensagem Nº {dados_frase['numero']} de {dados_frase['total']}
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    st.write("")
    st.markdown("---")
    
    # Botão Real de Download
    st.download_button(
        label=f"📥 Confirmar Download do {tipo_arquivo}",
        data=dados_arquivo,
        file_name=nome_arquivo,
        mime=mime_type,
        type="primary"
    )

# --- DOWNLOADS (ACIONADORES ATUALIZADOS) ---
st.markdown("---")
st.subheader("🚀 Baixar Arquivos")

nome_safe = st.session_state['dados']['nome'].strip() or "Parecer"
nome_sugestao = "".join([c for c in nome_safe if c.isalnum() or c in (' ','-','_')]).strip()

nome_final_arquivo = st.text_input("Nome do Arquivo para salvar (Word/PDF):", value=nome_sugestao)
if not nome_final_arquivo: nome_final_arquivo = "Parecer_Tecnico"

col_d1, col_d2 = st.columns(2)

# Lógica para WORD
with col_d1:
    if st.session_state['dados']['car']:
        # Pré-gera o arquivo Word
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Cm(3); sec.bottom_margin = Cm(2); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
        
        header_section = sec.header
        header_paragraph = header_section.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_paragraph.paragraph_format.left_indent = Cm(-1.75)
        
        if os.path.exists(ARQUIVO_LOGO):
             run_logo = header_paragraph.add_run()
             run_logo.add_picture(ARQUIVO_LOGO, width=Cm(1.59), height=Cm(1.59))

        # --- INSERÇÃO DA NUMERAÇÃO DE PÁGINA NO RODAPÉ WORD ---
        section = doc.sections[0]
        footer = section.footer
        footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
        
        run_footer = footer_p.add_run()
        criar_elemento_pagina(run_footer)
        # ------------------------------------------------------

        style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(12)
        
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Justificativa do Parecer Técnico"); r.bold = True; r.font.size = Pt(14)
        p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(f"CAR: {st.session_state['dados']['car']}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
        p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(f"SP-NOT: {st.session_state['dados']['sp_not']}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].bold = True
        
        doc.add_paragraph()
        
        p = doc.add_paragraph(); p.add_run("Nome do Imóvel Rural: ").bold = True; p.add_run(st.session_state['dados']['imovel'])
        p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(); p.add_run("Nome: ").bold = True; p.add_run(st.session_state['dados']['nome'])
        p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(); p.add_run("CPF/CNPJ: ").bold = True; p.add_run(formatar_documento(st.session_state['dados']['doc']))
        doc.add_paragraph()
        
        for i, item_obj in enumerate(st.session_state['dados']['itens']):
             titulo = item_obj['titulo']
             if not titulo: titulo = "Item Sem Título"
             id_item = item_obj['id']
             
             p = doc.add_paragraph(f"{i+1}. {titulo}"); p.runs[0].bold = True
             
             texto_item = st.session_state['dados']['textos'].get(id_item, "")
             widget_val = st.session_state.get(f"txt_area_{id_item}")
             if widget_val is not None: texto_item = widget_val
             
             linhas = texto_item.split('\n')
             
             for linha in linhas:
                 if not linha:
                     doc.add_paragraph()
                     continue
                 
                 p = doc.add_paragraph()
                 p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                 p.paragraph_format.first_line_indent = Cm(1.25)
                 
                 segmentos = parse_styled_text(linha)
                 for seg in segmentos:
                     run = p.add_run(seg['text'])
                     run.bold = seg['bold']
                     run.italic = seg['italic']
                     run.underline = seg['underline']

             lista_imgs = st.session_state['dados']['imagens_b64'].get(id_item, [])
             for b64_img in lista_imgs:
                 temp_path = b64_para_tempfile(b64_img)
                 if temp_path:
                     try:
                         p_img = doc.add_paragraph()
                         p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                         run_img = p_img.add_run()
                         run_img.add_picture(temp_path, width=Cm(7))
                     except Exception: pass
                     finally:
                         if os.path.exists(temp_path): os.unlink(temp_path)

             doc.add_paragraph("\n\n")
        
        p = doc.add_paragraph("________________________________________"); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(st.session_state['dados']['nome']); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)

        cidade_doc = st.session_state['dados']['cidade'] if st.session_state['dados']['cidade'] else "Mogi das Cruzes"
        p = doc.add_paragraph(f"{cidade_doc}, {obter_data_extenso()}."); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        buffer_word = io.BytesIO(); doc.save(buffer_word); buffer_word.seek(0)
        
        # BOTÃO PARA ACIONAR O MODAL (WORD)
        if st.button("⬇️ Gerar Word (.docx)"):
           gerar_nova_frase_motivacional() # GERA A FRASE E CONGELA ELA ANTES DE ABRIR
           mostrar_mensagem_download(
               "Word",
               buffer_word,
               f"{nome_final_arquivo}.docx",
               "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
           )

# Lógica para PDF
with col_d2:
    if pdf_bytes_preview:
        # BOTÃO PARA ACIONAR O MODAL (PDF)
        if st.button("⬇️ Gerar PDF (.pdf)"):
           gerar_nova_frase_motivacional() # GERA A FRASE E CONGELA ELA ANTES DE ABRIR
           mostrar_mensagem_download(
               "PDF",
               pdf_bytes_preview,
               f"{nome_final_arquivo}.pdf",
               "application/pdf"
           )

